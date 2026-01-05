"""
文件处理服务
负责处理各种格式的规则书文件，提取文本内容和元数据
"""

import os
import uuid
import tempfile
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from datetime import datetime

# 文件处理库
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

import markdown

from ..core.config import settings
from ..core.exceptions import StoryMasterValidationError
from ..core.logging import app_logger
from ..models.parsing_models import ProcessedFile


class FileProcessor:
    """文件处理器"""
    
    def __init__(self, upload_dir: str = None):
        self.upload_dir = Path(upload_dir or settings.upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.temp_dir = Path(tempfile.gettempdir()) / "storymaster_rulebooks"
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        
        # 支持的文件处理器
        self.processors = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'txt': self._process_txt,
            'json': self._process_json,
            'md': self._process_markdown
        }
    
    async def validate_file(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        验证文件格式和安全性
        
        Returns:
            Dict: 包含验证结果和元数据
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise StoryMasterValidationError(f"文件不存在: {file_path}")
        
        # 检查文件大小
        file_size = file_path.stat().st_size / (1024 * 1024)  # 转换为MB
        max_size = getattr(settings, 'max_rulebook_file_size', 50)
        if file_size > max_size:
            raise StoryMasterValidationError(
                f"文件大小超出限制。最大允许: {max_size}MB，当前: {file_size:.2f}MB"
            )
        
        # 检查文件类型
        if file_type.lower() not in self.processors:
            raise StoryMasterValidationError(
                f"不支持的文件类型: {file_type}。支持的类型: {', '.join(self.processors.keys())}"
            )
        
        # 生成文件哈希
        file_hash = self._generate_file_hash(file_path)
        
        return {
            "valid": True,
            "file_size": file_size,
            "file_hash": file_hash,
            "mime_type": self._get_mime_type(file_path)
        }
    
    async def process_uploaded_file(
        self, 
        file_content: bytes, 
        file_name: str, 
        file_type: str
    ) -> ProcessedFile:
        """
        处理上传的文件
        
        Args:
            file_content: 文件二进制内容
            file_name: 原始文件名
            file_type: 文件类型
            
        Returns:
            ProcessedFile: 处理后的文件对象
        """
        try:
            # 生成唯一文件名
            unique_id = uuid.uuid4().hex
            safe_filename = f"{unique_id}_{file_name}"
            file_path = self.upload_dir / safe_filename
            
            # 保存文件
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            # 验证文件
            validation_result = await self.validate_file(str(file_path), file_type)
            if not validation_result["valid"]:
                file_path.unlink()  # 删除无效文件
                raise StoryMasterValidationError(f"文件验证失败: {validation_result.get('error', '未知错误')}")
            
            # 处理文件内容
            processor = self.processors.get(file_type.lower())
            if not processor:
                raise StoryMasterValidationError(f"不支持的文件类型: {file_type}")
            
            content, metadata = await processor(str(file_path))
            
            # 内容分块
            content_chunks = await self._chunk_content(content)
            
            # 创建处理后的文件对象
            processed_file = ProcessedFile(
                file_path=str(file_path),
                file_name=file_name,
                file_type=file_type,
                file_size=len(file_content),
                content=content,
                content_chunks=content_chunks,
                metadata={
                    **metadata,
                    **validation_result,
                    "processed_at": datetime.now().isoformat()
                }
            )
            
            app_logger.info(f"文件处理成功: {file_name}, 大小: {len(file_content)/1024/1024:.2f}MB, 类型: {file_type}")
            
            return processed_file
            
        except Exception as e:
            app_logger.error(f"文件处理失败: {file_name}, 错误: {e}", exc_info=True)
            # 清理可能创建的文件
            if 'file_path' in locals() and file_path.exists():
                file_path.unlink()
            raise StoryMasterValidationError(f"文件处理失败: {str(e)}")
    
    async def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """处理PDF文件"""
        if not pdfplumber:
            raise StoryMasterValidationError("PDF处理库未安装，请安装pdfplumber")
        
        content_parts = []
        metadata = {
            "page_count": 0,
            "has_images": False,
            "title": "",
            "author": ""
        }
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)
                
                # 尝试提取PDF元数据
                if hasattr(pdf, 'metadata') and pdf.metadata:
                    metadata["title"] = pdf.metadata.get('Title', '')
                    metadata["author"] = pdf.metadata.get('Author', '')
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            content_parts.append(page_text)
                        
                        # 检查页面是否有图片
                        if hasattr(page, 'images') and page.images:
                            metadata["has_images"] = True
                            
                    except Exception as e:
                        app_logger.warning(f"处理PDF第{page_num}页失败: {e}")
                        continue
            
            content = "\n\n".join(content_parts)
            
            # 如果pdfplumber失败，尝试PyPDF2作为后备
            if not content.strip() and PyPDF2:
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        content = ""
                        for page in pdf_reader.pages:
                            content += page.extract_text() + "\n\n"
                except Exception as e:
                    app_logger.warning(f"使用PyPDF2处理PDF失败: {e}")
            
            return content, metadata
            
        except Exception as e:
            app_logger.error(f"PDF文件处理失败: {file_path}, 错误: {e}")
            raise StoryMasterValidationError(f"PDF文件处理失败: {str(e)}")
    
    async def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """处理Word文档"""
        if not Document:
            raise StoryMasterValidationError("Word处理库未安装，请安装python-docx")
        
        try:
            doc = Document(file_path)
            
            # 提取段落文本
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            content = "\n\n".join(paragraphs)
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "tables": tables
            }
            
            return content, metadata
            
        except Exception as e:
            app_logger.error(f"Word文件处理失败: {file_path}, 错误: {e}")
            raise StoryMasterValidationError(f"Word文件处理失败: {str(e)}")
    
    async def _process_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """处理文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # 尝试检测其他编码
            if not content.strip():
                encodings = ['gbk', 'gb2312', 'big5', 'latin-1']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            content = file.read()
                        if content.strip():
                            break
                    except UnicodeDecodeError:
                        continue
            
            line_count = len(content.splitlines())
            word_count = len(content.split())
            
            metadata = {
                "line_count": line_count,
                "word_count": word_count,
                "character_count": len(content),
                "encoding": "utf-8"
            }
            
            return content, metadata
            
        except Exception as e:
            app_logger.error(f"文本文件处理失败: {file_path}, 错误: {e}")
            raise StoryMasterValidationError(f"文本文件处理失败: {str(e)}")
    
    async def _process_json(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """处理JSON文件"""
        import json
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                json_data = json.load(file)
            
            # 如果JSON已经是规则书格式，直接返回
            if isinstance(json_data, dict) and 'entities' in json_data:
                content = json.dumps(json_data, indent=2, ensure_ascii=False)
                metadata = {
                    "format": "rulebook_schema",
                    "entity_count": len(json_data.get('entities', {})),
                    "rule_count": len(json_data.get('rules', {})),
                    "is_preformatted": True
                }
            else:
                # 否则转换为文本
                content = json.dumps(json_data, indent=2, ensure_ascii=False)
                metadata = {
                    "format": "json",
                    "is_preformatted": False
                }
            
            return content, metadata
            
        except json.JSONDecodeError as e:
            app_logger.error(f"JSON文件解析失败: {file_path}, 错误: {e}")
            raise StoryMasterValidationError(f"JSON文件格式错误: {str(e)}")
        except Exception as e:
            app_logger.error(f"JSON文件处理失败: {file_path}, 错误: {e}")
            raise StoryMasterValidationError(f"JSON文件处理失败: {str(e)}")
    
    async def _process_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """处理Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                markdown_content = file.read()
            
            # 提取原始Markdown内容
            content = markdown_content
            
            # 转换为HTML以获取结构信息
            html_content = markdown.markdown(markdown_content)
            
            # 统计信息
            lines = markdown_content.splitlines()
            headings = [line for line in lines if line.startswith('#')]
            code_blocks = markdown_content.count('```')
            
            metadata = {
                "format": "markdown",
                "line_count": len(lines),
                "heading_count": len(headings),
                "code_block_count": code_blocks // 2,
                "html_preview": html_content[:1000] + "..." if len(html_content) > 1000 else html_content
            }
            
            return content, metadata
            
        except Exception as e:
            app_logger.error(f"Markdown文件处理失败: {file_path}, 错误: {e}")
            raise StoryMasterValidationError(f"Markdown文件处理失败: {str(e)}")
    
    async def _chunk_content(self, content: str, chunk_size: int = 4000) -> List[str]:
        """将内容分块"""
        if not content:
            return []
        
        chunks = []
        paragraphs = content.split('\n\n')
        current_chunk = ""
        
        for paragraph in paragraphs:
            # 如果当前块加上新段落不超过限制，添加到当前块
            if len(current_chunk) + len(paragraph) + 2 <= chunk_size:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
            else:
                # 保存当前块
                if current_chunk:
                    chunks.append(current_chunk)
                
                # 如果单个段落就超过限制，强制分割
                if len(paragraph) > chunk_size:
                    for i in range(0, len(paragraph), chunk_size):
                        chunks.append(paragraph[i:i+chunk_size])
                    current_chunk = ""
                else:
                    current_chunk = paragraph
        
        # 添加最后一个块
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _generate_file_hash(self, file_path: str) -> str:
        """生成文件哈希值"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _get_mime_type(self, file_path: str) -> str:
        """获取文件MIME类型"""
        import mimetypes
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type or "application/octet-stream"